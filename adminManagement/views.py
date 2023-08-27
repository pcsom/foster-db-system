import decimal
from json.encoder import JSONEncoder
from django.contrib.messages.api import error
from django.core.exceptions import ValidationError, ObjectDoesNotExist
from django.db.models.fields import IntegerField
from django.db.models import Count
from django.db.utils import IntegrityError
from django.http.response import HttpResponse
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.models import User, auth, Group
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from decorators import unauthenticated_user, allowed_users
from accounts.models import child, userBasic, careType
from .models import *
from .forms import *
from django.conf import settings
from django.core.mail import send_mail
from utility import *
from django.http import JsonResponse
import datetime
from irequests.models import itemsRequest
import docx
from docx.shared import Cm, Pt
import io
import json
from django.db import transaction
import ast
from accounts.forms import userBasicForm
import bisect
import time
import logging
import traceback
import sys
from django.contrib.auth import password_validation
log = logging.getLogger(__name__)

@login_required
@allowed_users(allowed_roles=['admin',])
def getFamilyDataArray(request):
    #possibly use timeit in the future to compare if two queries or one query plus list comprehension to find dataTableInfo's is faster
    for i in userBasic.objects.all():
        i.save()

    data2 = userBasic.objects.values_list('dataTableInfo')
    dataList = []
    infoList = eval(request.POST['infoList'])

    for i in data2:
        curDict = i[0]
        if curDict['result'][0] == "None":
            noneDict = {'edit': False, 'result':["None",] + ['--' for x in range(1, len(infoList))], 'id':curDict['id']}    
            noneDict['result'][infoList["# of Children"]] = curDict['numChild']
            noneDict['result'][infoList["# of Requests"]] = curDict['numReqs']
            dataList.append(noneDict)
        else:
            dataList.append(curDict)
    return JsonResponse(dataList, safe=False)



@login_required
@allowed_users(allowed_roles=['admin',])
def updateFamily(request):
    objId = request.POST.get('id')
    result = request.POST.getlist('result[]')
    test = User.objects.get(id=objId)
    alertList = []
    save, nameChange = True, False

    test.email = str(result[3])

    nameChange = test.username != str(result[0])
    test.username = str(result[0])

    test.first_name = str(result[1])
    test.last_name = str(result[2])
    test.userbasic.address = str(result[4])
    test.userbasic.city = str(result[5])
    test.userbasic.state = str(result[6])
    test.userbasic.zipCode = int(result[7])
    test.userbasic.country = str(result[8])
    test.userbasic.cellPhone = str(result[9])
    test.userbasic.homePhone = str(result[10])

    newCont = str(result[11])
    if newCont == 'Cell Phone Number':
        test.userbasic.contactMethod = 'cellPhRadio'
    elif newCont == 'Home Phone Number':
        test.userbasic.contactMethod = 'homePhRadio'
    else:
        test.userbasic.contactMethod = 'emailRadio'

    test.userbasic.emailUpdateList = str(result[12]) == 'true'
    test.userbasic.remindUpdateList = str(result[13]) == 'true'

    try:
        with transaction.atomic():
            test.userbasic.save()
            test.save()
    except Exception as e:
        if type(e).__name__ == "ValidationError":
            theDict = ast.literal_eval(str(e))
            for singleField in theDict.keys():
                for mess in theDict[singleField]:
                    alertList.append(f"Field '{singleField}' - {mess}")
        elif type(e).__name__ == "ValueError":
            mess = str(e)
            field = mess[7:(mess[7:].find("'") + 7)]
            alertList.append(mess.replace(field, processName(field)))
        else: 
            alertList.append(str(e))
    else:
        if nameChange:
            subject = 'NAFC Username Changed'
            message = f"Hello {test.first_name} {test.last_name},\n\nYour North Alabama Foster Closet username was changed. Your new username is '{test.username}'\nPlease remember to use this new username when logging into the NAFC website. Thank you!\n\nNAFC Team"
            recipient_list = [str(test.email),]
            
            GMAIL_SERVICE = initGmailAPI()
            sent = send_message(GMAIL_SERVICE, 'me', create_message("me", ", ".join(recipient_list), subject, message))
            if not sent:
                messages.error(request, "Was not able to notify this user about their new username via email. Please ensure the email system is working.")


    return JsonResponse(alertList, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def allFamilyInfo(request):
    fieldsList = ['username', 'firstName', 'lastName', 'email']
    #fieldsList.extend(list(to_dict(userBasic.objects.first(), exclude=['user', 'registrationProgress']).keys()))
    fieldsList.extend(getKeys(userBasic.objects.first(), exclude={'id', 'user', 'registrationProgress'}))
    beforeLen = len(fieldsList)
    nonEdit = ["# of Children", "# of Types of Care", "Adoptive Family?", "Foster Family?", "Kinship Family?", "Safety Plan Family?", "# of Requests"]
    
    fieldsList.extend(nonEdit)
    infoList = {}
    for x in range(0, len(fieldsList)):
        infoList[processName(fieldsList[x])] = x

    context = {
        #'infoListOld': [{"num": x, "field": processName(fieldsList[x])} for x in range(0, len(fieldsList))], 
        'infoList': infoList, 
        'nonEdit': json.dumps([x + beforeLen for x in range(0, len(nonEdit))]),
    }

    return render(request, 'admin/familyInfo.html', context)





@login_required
@allowed_users(allowed_roles=['admin',])
def viewFamily(request, pk):
    toView = User.objects.get(id=int(pk))
    st = toView.is_staff
    basicInfo = toView.userbasic

    if request.method == "POST":
        mess = "Confirmation for changing user permissions failed."
        if 'confirm' in request.POST:
            if(st):
                toView.is_staff = False
                toView.save()
                mess = "This user's admin permissions have been revoked."
            else:
                toView.is_staff = True
                toView.save()
                mess = "This user has been made an admin."
        messages.success(request, mess)
        return redirect(f'/view-family/{pk}')


    children = child.objects.filter(parent=toView).all()
    currentCareType = careType.objects.filter(user=toView).first()
    
    context = {
        'children': children, 'basicInfo': basicInfo, 'curUser': toView, 'pn': str(toView) + "'s",
        'currentCareType': currentCareType, 'st':st, 'regProg':toView.userbasic.registrationProgress == "Complete"
    }

    return render(request, 'admin/viewSingleFamily.html', context)





@login_required
@allowed_users(allowed_roles=['admin',])
def changeRegProg(request, pk):
    toManage = userBasic.objects.get(user=int(pk))
    progDone = toManage.registrationProgress == "Complete"
    print(request.POST)
    mess = "Confirmation for changing user's registration progress failed."
    if 'confirm' in request.POST:
        if progDone:
            toManage.registrationProgress = "type-of-care"
            toManage.save()
            mess = "User has been required to verify care type information and re-sign Statement of Truthfulness"
        else:
            toManage.registrationProgress = "Complete"
            toManage.save()
            mess = "User has been allowed to bypass care type selection and Statement of Truthfulness"
    messages.success(request, mess)
    return redirect(f'/view-family/{pk}')






@login_required
@allowed_users(allowed_roles=['admin',])
def allChildInfo(request):
    context = {}
    fieldsList = []
    try:
        fieldsList = getKeys(child.objects.first(), exclude=set({'id','itemsrequest','distributionentry'})) + ['# of Requests',]
    except Exception as e:
        log.debug(f"{traceback.format_exc()}")
        log.debug("in the error")
        context = {'noChild': True}
    else:
        infoList = {}
        for x in range(0, len(fieldsList)):
            infoList[processName(fieldsList[x])] = x
        context = {'infoList': infoList, 'noChild': False}
    log.debug(json.dumps(context))
    return render(request, 'admin/allChildInfo.html', context)

    



@login_required
@allowed_users(allowed_roles=['admin',])
def getChildDataArray(request):    
    for i in child.objects.all():
        i.save()
    return JsonResponse(list(child.objects.values_list('dataTableInfo', flat=True)), safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def addDistLogEntry(request):
    if request.method == 'POST':
        alertList = []
        form = distributionForm(request.POST)
        
        if form.is_valid():
            try:
                with transaction.atomic():
                    newEntry = formToDistribModel(form)
                    curChild = child.objects.get(id=request.POST["child"])
                    isSent = 'sent' in request.POST
                    for aReq in itemsRequest.objects.filter(forChild=curChild, status="Receiving"):
                        if isSent:
                            aReq.status = "Fulfilled"
                            aReq.save()
                    newEntry.forChild = curChild
                    newEntry.sent = isSent
                    newEntry.save()
            except Exception as e:
                if type(e).__name__ == "ValidationError":
                    theDict = ast.literal_eval(str(e))
                    for singleField in theDict.keys():
                        for mess in theDict[singleField]:
                            alertList.append(f"Field '{singleField}' - {mess}")
                elif type(e).__name__ == "ValueError":
                    mess = str(e)
                    field = mess[7:(mess[7:].find("'") + 7)]
                    alertList.append(mess.replace(field, processName(field)))
                else: 
                    alertList.append(str(e))
            else:
                messages.success(request, "Succesfully added distribution log entry.")
        else:
            for singleField in form.errors.keys():
                for singleMess in form.errors[singleField]:
                    alertList.append(f"{singleField} - {singleMess}")

        return JsonResponse(alertList, safe=False)
    else:
        form = distributionForm()
        context = {'form': form, 'userList': [(x.id, str(x.userbasic)) for x in User.objects.all()]}
        return render(request, "admin/addDistLogEntry.html", context)



@login_required
@allowed_users(allowed_roles=['admin',])
def getUserChildren(request):
    return JsonResponse([(x.id, str(x)) for x in child.objects.filter(parent=request.POST['userId'])], safe=False)



@login_required
@allowed_users(allowed_roles=['admin',])
def viewDistLog(request):
    fieldsList = ['For Family',]
    try:
        fieldsList.extend(getKeys(distributionEntry.objects.first(), exclude=set({'id',}), overrideName={'sent':'Have Items Been Sent?'}))
    except:
        messages.error(request, "There are currently no distribution log entries!")
        return redirect('/')

    disp = ReqFormSettings.load().databaseDisplay
    fieldSet = set()
    for i in distributionEntry.objects.all():
        for j in json.loads(i.formData):
            fieldSet.add((j, disp[j]))
    
    curFieldsTup = sorted(list(fieldSet), key=lambda a: a[1])
    curFields = [x[1] for x in curFieldsTup]
    curDbns = [x[0] for x in curFieldsTup]
    
    context = {'infoList': [{"num": x, "field": processName(fieldsList[x])} for x in range(0, len(fieldsList))], 'curFields':curFields, 'curDbns':curDbns}
    return render(request, 'admin/distLog.html', context)




@login_required
@allowed_users(allowed_roles=['admin',])
def getDistDataArray(request):
    users = {}
    for i in distributionEntry.objects.all():
        users[i.pk] = str(i.forChild.parent)
        i.save()

    data2 = list(distributionEntry.objects.values_list('dataTableInfo', flat=True))
    dataList = [data2, users, [(x.id, str(x)) for x in child.objects.all()]]

    return JsonResponse(dataList, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def updateDist(request, pk):
    test = distributionEntry.objects.get(id=pk)


    if request.method == 'POST':
        form = distributionForm(request.POST)
        alertList = []
        if form.is_valid():
            try:
                with transaction.atomic():
                    newEntry = formToDistribModel(form)
                    test.dateDistributed = newEntry.dateDistributed
                    test.formData = newEntry.formData
                    curChild = child.objects.get(id=request.POST["child"])
                    isSent = 'sent' in request.POST
                    for aReq in itemsRequest.objects.filter(forChild=curChild, status="Receiving"):
                        if isSent:
                            aReq.status = "Fulfilled"
                            aReq.save()
                    test.forChild = curChild
                    test.sent = isSent
                    test.save()
            except Exception as e:
                if type(e).__name__ == "ValidationError":
                    theDict = ast.literal_eval(str(e))
                    for singleField in theDict.keys():
                        for mess in theDict[singleField]:
                            alertList.append(f"Field '{singleField}' - {mess}")
                elif type(e).__name__ == "ValueError":
                    mess = str(e)
                    field = mess[7:(mess[7:].find("'") + 7)]
                    alertList.append(mess.replace(field, processName(field)))
                else: 
                    alertList.append(str(e))
            else:
                messages.success(request, "Succesfully added distribution log entry.")
        else:
            for singleField in form.errors.keys():
                for singleMess in form.errors[singleField]:
                    alertList.append(f"{singleField} - {singleMess}")
        return JsonResponse(alertList, safe=False)

    else:
        uform = distributionForm()
        uform.fields['dateDistributed'].initial = test.dateDistributed.strftime("%m/%d/%Y")
        formd = json.loads(test.formData)
        for i in formd:
            uform.fields[i].initial = formd[i]
        print(test.forChild.parent.id)
        context = {'form': uform, 'pk':pk, 'userList': [(x.id, str(x.userbasic)) for x in User.objects.all()], 'test':test}
        return render(request, 'admin/updateDist.html', context)


"""def updateDist(request):
    objId = request.POST.get('id')
    result = request.POST.getlist('result[]')
    test = distributionEntry.objects.get(id=objId)
    fieldsList = getKeys(test, exclude=set({'id',}))
    alertList = []
    month, day, year = 0, 0, 0

    if result[0].isdigit():
        test.forChild = child.objects.get(id=result[0])

    try:
        month, day, year = map(int, result[1].split('/'))
    except ValueError as e:
        alertList.append('Invalid date provided.')
        return JsonResponse(alertList, safe=False)
    test.dateDistributed = datetime.date(year, month, day)

    if result[2] == "Yes":      #No need to check for "No" as there will be no option to go from Yes to No
        test.sent = True
        for i in itemsRequest.objects.filter(forChild=test.forChild, status="Receiving"):
            i.status = "Fulfilled"
            i.save()

    result = result[3:]


    for x in range(3, len(fieldsList)):
        #exec(f"test.{fieldsList[x]} = '{result[x-2]}'")
        setattr(test, fieldsList[x], result[x-3])
        #x += 1

    try:
        with transaction.atomic():
            test.save()
    except Exception as e:
        if type(e).__name__ == "ValidationError":
            theDict = ast.literal_eval(str(e))
            for singleField in theDict.keys():
                for mess in theDict[singleField]:
                    alertList.append(f"Field '{singleField}' - {mess}")
        elif type(e).__name__ == "ValueError":
            mess = str(e)
            field = mess[7:(mess[7:].find("'") + 7)]
            alertList.append(mess.replace(field, processName(field)))
        else: 
            alertList.append(str(e))
    return JsonResponse(alertList, safe=False)"""




@login_required
@allowed_users(allowed_roles=['admin',])
def deleteDist(request):
    entry = distributionEntry.objects.get(id=int(request.POST['id']))
    entry.delete()

    return JsonResponse("Deleted!", safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def addDonorLogEntry(request):
    if request.method == 'POST':
        alertList = []
        form = donorForm(request.POST)
        if form.is_valid():
            try:
                with transaction.atomic():
                    newEntry = form.save()
                    newEntry.save()
            except Exception as e:
                if type(e).__name__ == "ValidationError":
                    theDict = ast.literal_eval(str(e))
                    for singleField in theDict.keys():
                        for mess in theDict[singleField]:
                            alertList.append(f"Field '{singleField}' - {mess}")
                elif type(e).__name__ == "ValueError":
                    mess = str(e)
                    field = mess[7:(mess[7:].find("'") + 7)]
                    alertList.append(mess.replace(field, processName(field)))
                else: 
                    alertList.append(str(e))
            else:
                messages.success(request, "Successfully added donor log entry.")
        else:
            for singleField in form.errors.keys():
                for singleMess in form.errors[singleField]:
                    alertList.append(f"{singleField} - {singleMess}")
        return JsonResponse(alertList, safe=False)
    else:
        form = donorForm()
        context = {'form': form}
        return render(request, "admin/addDonorLogEntry.html", context)




@login_required
@allowed_users(allowed_roles=['admin',])
def viewDonorLog(request):
    fieldsList = []
    try:
        fieldsList = getKeys(donorEntry.objects.first(), exclude=set({'id',}))
    except:
        messages.error(request, "There are currently no donor log entries!")
        return redirect('/')
    context = {'infoList': [{"num": x, "field": processName(fieldsList[x])} for x in range(0, len(fieldsList))],}
    return render(request, 'admin/donorLog.html', context)




@login_required
@allowed_users(allowed_roles=['admin',])
def getDonorDataArray(request):    
    for i in donorEntry.objects.all():
        i.save()

    return JsonResponse(list(donorEntry.objects.values_list('dataTableInfo', flat=True)), safe=False)



@login_required
@allowed_users(allowed_roles=['admin',])
def updateDonor(request):
    objId = request.POST.get('id')
    result = request.POST.getlist('result[]')
    test = donorEntry.objects.get(id=objId)
    fieldsList = getKeys(test, exclude=set({'id',}))
    alertList = []

    for x in range(0, len(fieldsList)):
        try:
            exec(f"test.{fieldsList[x]} = '{result[x]}'")
        except SyntaxError:
            cur = result[x].replace('\n', '\\n')
            exec(f"test.{fieldsList[x]} = '{cur}'")
        x += 1

    try:
        with transaction.atomic():
            test.save()
    except Exception as e:
        if type(e).__name__ == "ValidationError":
            theDict = ast.literal_eval(str(e))
            for singleField in theDict.keys():
                for mess in theDict[singleField]:
                    alertList.append(f"Field '{singleField}' - {mess}")
        elif type(e).__name__ == "ValueError":
            mess = str(e)
            field = mess[7:(mess[7:].find("'") + 7)]
            alertList.append(mess.replace(field, processName(field)))
        else: 
            alertList.append(str(e))
    return JsonResponse(alertList, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def deleteDonor(request):
    entry = donorEntry.objects.get(id=int(request.POST['id']))
    entry.delete()

    return JsonResponse("Deleted!", safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def addHoursLogEntry(request):
    if request.method == 'POST':
        alertList = []
        volId = request.POST['selVol']
        volObj = None
                
        
        form = hoursForm(request.POST)
        if form.is_valid():
            try:
                with transaction.atomic():
                    if not volId.isdigit():
                        volObj = volunteer.objects.create(name=volId, totalHoursWorked=0)
                    else:
                        volObj = volunteer.objects.get(id=volId)

                    
                    newEntry = form.save(commit=False)
                    newEntry.forVolunteer = volObj
                    newEntry.save()

                    volObj.totalHoursWorked += newEntry.hoursWorked
                    volObj.save()

                    globObj = volunteer.objects.get(name="Global Volunteer Hour Count")
                    globObj.totalHoursWorked += newEntry.hoursWorked
                    globObj.save()
            except Exception as e:
                if type(e).__name__ == "ValidationError":
                    theDict = ast.literal_eval(str(e))
                    for singleField in theDict.keys():
                        for mess in theDict[singleField]:
                            alertList.append(f"Field '{singleField}' - {mess}")
                elif type(e).__name__ == "ValueError":
                    mess = str(e)
                    field = mess[7:(mess[7:].find("'") + 7)]
                    alertList.append(mess.replace(field, processName(field)))
                else: 
                    alertList.append("SYSTEM BACKEND ERROR: " + str(e))
            else:
                #volObj.updateVolHours()
                #updateGlobalVolHours()
                #rectifyGlobalVolHours()

                messages.success(request, "Successfully added hours log entry.")
        else:
            for singleField in form.errors.keys():
                for singleMess in form.errors[singleField]:
                    alertList.append(f"{singleField} - {singleMess}")
        return JsonResponse(alertList, safe=False)
    else:
        form = hoursForm()
        context = {'form': form, 'vols': [(vol.id, vol.name) for vol in volunteer.objects.all()]}
        return render(request, "admin/addHoursLogEntry.html", context)




@login_required
@allowed_users(allowed_roles=['admin',])
def viewHoursLog(request):
    fieldsList = []

    try:
        fieldsList = getKeys(hoursEntry.objects.first(), exclude=set({'id',}))
    except:
        messages.error(request, "There are currently no volunteer hours log entries!")
        return redirect('/')
    
    context = {'infoList': [{"num": x, "field": processName(fieldsList[x])} for x in range(0, len(fieldsList))],}
    return render(request, 'admin/hoursLog.html', context)




@login_required
@allowed_users(allowed_roles=['admin',])
def getHoursDataArray(request):
    for i in hoursEntry.objects.all():
        i.save()

    data2 = list(hoursEntry.objects.values_list('dataTableInfo', flat=True))
    dataList = [data2, list(volunteer.objects.values_list('id', 'name'))]
    return JsonResponse(dataList, safe=False)



@login_required
@allowed_users(allowed_roles=['admin',])
def updateHours(request):
    objId = request.POST.get('id')
    result = request.POST.getlist('result[]')
    test = hoursEntry.objects.get(id=objId)
    alertList = []
    month, day, year = 0, 0, 0

    try:
        month, day, year = map(int, result[0].split('/'))
    except ValueError as e:
        alertList.append('Invalid date provided.')
        return JsonResponse(alertList, safe=False)
    test.date = datetime.date(year, month, day)

    if result[1].isdigit():           #The javascript returns a number for result[1] if the volunteer field had changed
        test.forVolunteer = volunteer.objects.get(id=result[1])
    
    prevHours = test.hoursWorked
    test.hoursWorked = decimal.Decimal(result[2])

    try:
        with transaction.atomic():
            test.save()

            if request.POST['self'] == 'true':
                volObj = test.forVolunteer
                volObj.totalHoursWorked -= prevHours
                volObj.totalHoursWorked += test.hoursWorked
                volObj.save()

                globObj = volunteer.objects.get(name="Global Volunteer Hour Count")
                if volObj != globObj:
                    globObj.totalHoursWorked -= prevHours
                    globObj.totalHoursWorked += test.hoursWorked
                    globObj.save()
    except Exception as e:
        traceback.print_exc()
        if type(e).__name__ == "ValidationError":
            theDict = ast.literal_eval(str(e))
            for singleField in theDict.keys():
                for mess in theDict[singleField]:
                    alertList.append(f"Field '{singleField}' - {mess}")
        elif type(e).__name__ == "ValueError":
            mess = str(e)
            field = mess[7:(mess[7:].find("'") + 7)]
            alertList.append(mess.replace(field, processName(field)))
        else: 
            alertList.append(str(e))
    else:
        # if request.POST['self'] == 'true':
        #     test.forVolunteer.updateVolHours()
        print(request.POST.getlist('changes[]'))
        for volId in request.POST.getlist('changes[]'):         #changes array is for when the forVolunteer of an hourslog is changed
            volunteer.objects.get(id=volId).updateVolHours()        
        #rectifyGlobalVolHours()
        #updateGlobalVolHours()

    return JsonResponse(alertList, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def deleteHours(request):
    entry = hoursEntry.objects.get(id=int(request.POST['id']))
    prevHours = entry.hoursWorked
    volObj = entry.forVolunteer
    entry.delete()

    globObj = volunteer.objects.get(name="Global Volunteer Hour Count")
    globObj.totalHoursWorked -= prevHours
    globObj.save()

    volObj.totalHoursWorked -= prevHours
    volObj.save()


    #volObj.updateVolHours()
    #rectifyGlobalVolHours()
    #updateGlobalVolHours()

    return JsonResponse("Deleted!", safe=False)




# def getDate(a, forUser):
#     if forUser:
#         if len(a['data']) == 0:
#             return 999999999999
#         else: 
#             theDate = a['data'][0]['date']
#             return int(f"{theDate.year}{theDate.month}{theDate.day}")
#     else:
#         return int(f"{a['date'].year}{a['date'].month}{a['date'].day}")

def getDate(a, forUser):
    if forUser:
        if len(a['data']) == 0:
            return 999999999999
        else: 
            return a['data'][0]['date']
    else:
        return a['date']

@login_required
@allowed_users(allowed_roles=['admin',])
def viewAllRequests(request, pk):
    if request.method == "POST":
        dbId = request.POST['dbId']
        curReq = itemsRequest.objects.get(id=dbId)
        return JsonResponse(curReq.get_fields(strDate=True), safe=False)
    else:
        btns = []
        if pk == "11":
            btns = [{'desc': 'Show (Recieving) and (Fulfilled) requests too', 'link': '01'}, {'desc': 'Show the requests separately (remove grouping)', 'link': '10'}]
        elif pk == "10":
            btns = [{'desc': 'Show (Recieving) and (Fulfilled) requests too', 'link': '00'}, {'desc': 'Group the requests by user', 'link': '11'}]
        elif pk == "01":
            btns = [{'desc': "Hide (Recieving) and (Fulfilled) requests", 'link': '11'}, {'desc': 'Show the requests separately (remove grouping)', 'link': '00'}]
        elif pk == "00":
            btns = [{'desc': "Hide (Recieving) and (Fulfilled) requests", 'link': '10'}, {'desc': 'Group the requests by user', 'link': '01'}]
        else:
            return redirect('/view-all-requests/11')

        pendList = ['Pending', 'Collecting']

        if pk == "11" or pk == "01":
            users = User.objects.all()
            usersList = []

            for toManage in users:
                status = "(N/A)"
                pendCount = 0

                itemsRequestsRaw = itemsRequest.objects.filter(forUser=toManage).order_by("date")
                iRequests = []
                for one in itemsRequestsRaw:
                    if pk == "01" or pk == "11" and one.status in pendList:
                        iRequests.append({'id':len(iRequests), 'dbId':one.id, 'childDateInfo':one.childDateInfo(), 'date':str(one.date)})
                    if one.status in pendList:
                        pendCount += 1
                if len(itemsRequestsRaw) > 0:
                    for i in itemsRequestsRaw:
                        print(i.status)
                    print("end")
                    status = f"({itemsRequestsRaw.last().status})"

                if len(iRequests) == 0:
                    continue
                
                userInfo = {'name': f"{status} {str(toManage)}", 'data': iRequests}
                if pendCount > 0:
                    userInfo['button'] = 'Generate pending requests printout for this user'
                    userInfo['btnId'] = toManage.id
                else:
                    userInfo['button'] = 'None'

                usersList.append(userInfo)
                

            usersList.sort(key=lambda x: getDate(x, True))
            context = {'usersList': usersList, 'usersListJSON': json.dumps(usersList), 'group': "1", 'btns': btns, 'pk':pk}
            return render(request, 'admin/viewAllRequests.html', context)

        elif pk == "00" or pk == "10":
            itemsRequestsRaw = itemsRequest.objects.all()
            iRequests = []

            for one in itemsRequestsRaw:
                if pk == "00" or pk == "10" and one.status in pendList:
                    iRequests.append({'id':len(iRequests), 'dbId':one.id, 'dbId':one.id, 'name':str(one), 'date':str(one.date)})

            iRequests.sort(key=lambda x: getDate(x, False))

            context = {'iRequests': iRequests, 'iRequestsJSON': json.dumps(iRequests), 'group': "0", 'btns': btns, 'pk':pk}
            return render(request, 'admin/viewAllRequests.html', context)







def addFields(orig, add):
    for name, val in add:
        value = str(val)
        if name in orig:
            if orig[name] != value and value != ' ':
                orig[name] += ". " + value
        else:
            orig[name] = value
    return orig


@login_required
@allowed_users(allowed_roles=['admin',])
def getRequestReport(request, pk):                          #CHECK BACK ON EFFICIENCY OF THIS ONE
    toManage = User.objects.get(id=pk)
    itemsRequestsRaw = itemsRequest.objects.filter(forUser=toManage).all()

    iRequests = {}
    receivingListElems = ""
    for one in itemsRequestsRaw:
        if one.status == "Pending" or one.status == "Collecting":
            fieldsDict = one.get_fields()
            orig = {}
            try:
                orig = iRequests[str(one.forChild.id)]
            except KeyError:
                orig={}

            iRequests[str(one.forChild.id)] = addFields(orig, fieldsDict['fields'])
            if receivingListElems == "":
                receivingListElems = ", ".join(one.preferredMethodOfReceiving['info'])
                # for x in one.preferredMethodOfReceiving['info']:
                #     receivingListElems += x + ', '
            if one.status == "Pending":
                one.status = "Collecting"
                one.save()
        

    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    r1 = doc.add_paragraph()
    r2 = r1.add_run(f"Pending Requests for {toManage.first_name} {toManage.last_name}, {toManage.email}")
    r2.font.size = Pt(14)
    r1.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    pa = p.add_run(f"\t-Username: {toManage.username}")
    pa.font.size = Pt(10)

    basic = to_dict(toManage.userbasic, exclude=['user', 'registrationProgress', 'id'], display=['contactMethod',])
    bk = list(basic.keys())
    for x in range(0, len(bk), 2):
        theStr = f"\n\t-{processName(str(bk[x]))}: {basic[bk[x]]}"
        if x!=len(bk)-2:
            theStr+="\t"
        if x!=len(bk)-1:
            theStr += f"\t\t\t\t\t-{processName(str(bk[x+1]))}: {basic[bk[x+1]]}"
        px = p.add_run(theStr)
        px.font.size = Pt(10)

    
    if len(iRequests.keys()) != 0:
        pu = doc.add_paragraph()
        pu.paragraph_format.space_after = Pt(0)
        r4 = pu.add_run(f"\nPreferred method of receiving items:\n{receivingListElems}\n\n")
        r4.font.size = Pt(12)
        for childId in list(iRequests.keys()):
            theChild = child.objects.get(id=childId)
            
            para = f"Child name: {str(theChild).upper()}:"
            childInfo = to_dict(theChild, exclude=['parent', 'firstName', 'id'])
            # for x in childInfo:
            #     para += f"\n\t-{processName(str(x))}: {childInfo[x]}"

            
            ck = list(childInfo.keys())
            for x in range(0, len(ck), 2):
                if childInfo[ck[x]]:
                    curstr = f"\n\t-{processName(str(ck[x]))}: {childInfo[ck[x]]}"
                    if x!=len(ck)-1:
                        while len(curstr.expandtabs()) < 64:
                            curstr += "\t"
                        curstr += f"-{processName(str(ck[x+1]))}: {childInfo[ck[x+1]]}"
                    para += curstr


            
            para += "\n\n\t-Requested Items:"
            itemKeys, itemVals = list(iRequests[childId].keys()), list(iRequests[childId].values())
            for x in range(0, len(itemKeys)):
                para += f"\n\t\t-{itemKeys[x]}{itemVals[x]}"
            para += '\n'
            parad = doc.add_paragraph().add_run(para)
            parad.font.size = Pt(10)

    else:
        noPick = doc.add_paragraph().add_run("There are no requests for this user.")
        noPick.font.size = Pt(12)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    response = HttpResponse(bio.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    theStr = f"Requests for {toManage.first_name} {toManage.last_name}.docx"
    response["Content-Disposition"] = f'attachment; filename = "{theStr}"'
    response["Content-Encoding"] = "UTF-8"
    return response




@login_required
@allowed_users(allowed_roles=['admin',])
def getUserRequests(request):
    toManage = User.objects.get(id=request.POST.get('id'))
    noPend = True
    receiving = None

    iReqs = itemsRequest.objects.filter(forUser=toManage)       #Should make forUser an index. db_index=True 
    if len(iReqs) > 0:
        receiving = iReqs.first().preferredMethodOfReceiving['info']
        # if binarySearch(iReqs, 2, "{'Fulfilled':0, 'Receiving':1, 'Collecting':2, 'Pending':3}[v.status]") != -1:
        #     noPend = False
        if itemsRequest.objects.filter(forUser=toManage, status="Collecting"):                          #CHECK IF THIS WILL WORK W/O PROBS
            noPend = False
    else:
        receiving = ["ERROR: This user has never made a request.",]

    #Note that a query with filter takes O(log N) if B-tree indexed, else takes O(N). It also took longer in testing

    userDict = {'email': toManage.email, 'homePhone': toManage.userbasic.homePhone, 
        'cellPhone': toManage.userbasic.cellPhone, 
        'contactMethod': toManage.userbasic.get_contactMethod_display(), 
        'name': f"{toManage.first_name} {toManage.last_name}",
        'receiving': receiving, 'noPend': noPend, 'reqNames': [str(x)[len(str(toManage))+14:] for x in iReqs]}


    return JsonResponse(userDict, safe=False)




def sortReqs(aReq):
    return aReq

@login_required
@allowed_users(allowed_roles=['admin',])
def notifyRequests(request):
    if request.method == "POST":   
        uId = request.POST['userSelect']
        toManage = User.objects.get(id=uId)
        reqs = itemsRequest.objects.filter(forUser=toManage, status="Collecting").all() 
        for aReq in reqs:
            aReq.status = "Receiving"
            aReq.save()


        subject = 'NAFC Requested Items Delivery/Pickup'
        message = request.POST['theMessage']
        recipient_list = [toManage.email,]
        
        
        GMAIL_SERVICE = initGmailAPI()
        sent = send_message(GMAIL_SERVICE, 'me', create_message("me", ", ".join(recipient_list), subject, message))

        if sent:
            messages.success(request, f"Message has been sent to {str(toManage)}")
        else:
            messages.error(request, f"Message was not sent to {str(toManage)}, as the Gmail service is not responding.")


        return redirect("/")
    else:
        # userList = User.objects.all()
        # userIdName = []
        # for aUser in userList:
        #     if aUser.username != 'None':
        #         userIdName.append((aUser.id, str(aUser)))

        #Check if the following works without flaw
        userIdName = [(i[3], f"{i[0]} {i[1]} ({i[2]})") for i in itemsRequest.objects.filter(status="Collecting").values_list('forUser__first_name', 'forUser__last_name', 'forUser__username', 'forUser__id').distinct()]

        context = {'userList': userIdName, 'isEmpty': len(userIdName)==0}
        return render(request, 'admin/notifyRequests.html', context)


@login_required
@allowed_users(allowed_roles=['admin',])
def createUser(request):
    if request.method == 'POST':
        first_name = request.POST['first_name']
        last_name = request.POST['last_name']
        email_ = request.POST['email']
        username = request.POST['username']
        password1 = request.POST['password1']
        password2 = request.POST['password2']
        alertList = []
        pw = False

        if password1!=password2:
            alertList.append('Password confirmation failed! Please re-enter in the password fields.')
        elif User.objects.filter(username=username).exists():
            alertList.append('This username is taken!')
        elif User.objects.filter(email=email_).exists():
            alertList.append('There is an account already registered with this email!')
        else:
            form = userBasicForm(request.POST)
            if form.is_valid():
                try:
                    password_validation.validate_password(password1, request.user)
                    pw = True
                    with transaction.atomic():
                        newUserBasic = form.save()
                        newUser = User.objects.create_user(username=username, password=password1, email=email_, first_name=first_name, last_name=last_name)
                        if 'isAdmin' in request.POST:
                            newUser.is_staff = True
                            newUser.groups.add(Group.objects.get(name='admin'))

                        if 'bypassSign' in request.POST:
                            newUserBasic.registrationProgress='Complete'
                        else:
                            newUserBasic.registrationProgress='type-of-care'

                        newUser.groups.add(Group.objects.get(name='user'))
                        newUser.save()

                        newUserBasic.user = newUser
                        # if newUser.is_staff:
                        #     newUserBasic.registrationProgress='Complete'
                        #     careType.objects.create(user=newUser, name=newUser.first_name + " " + newUser.last_name + ", " + newUser.email)
                        # else:
                        #     newUserBasic.registrationProgress='type-of-care'
                        newUserBasic.save()
                except Exception as e:
                    if type(e).__name__ == "ValidationError":
                        theDict = ast.literal_eval(str(e))
                        if pw:
                            for singleField in theDict.keys():
                                for mess in theDict[singleField]:
                                    alertList.append(f"Field '{singleField}' - {mess}")
                        else:
                            alertList.extend([mess for mess in theDict])
                    elif type(e).__name__ == "ValueError":
                        mess = str(e)
                        field = mess[7:(mess[7:].find("'") + 7)]
                        alertList.append(mess.replace(field, processName(field)))
                    else: 
                        alertList.append(str(e))

                else:
                    messages.success(request, "New user successfully created.")
            else:
                for singleField in form.errors.keys():
                    for singleMess in form.errors[singleField]:
                        alertList.append(f"{singleField} - {singleMess}")
        return JsonResponse(alertList, safe=False)
    else:
        form = userBasicForm()
        context = {'form': form}
        return render(request, 'admin/createUser.html', context)



@login_required
@allowed_users(allowed_roles=['admin',])
def allVolInfo(request):
    fieldsList = []
    try:
        fieldsList = getKeys(volunteer.objects.first(), exclude=set({'id', 'hoursentry'}))
    except:
        messages.error(request, "There are currently no volunteers in the system!")
        return redirect('/')
    
    nonEdit = ['totalHoursWorked',]
    nonEditNum = [fieldsList.index(i) for i in nonEdit]

    context = {'infoList': [{"num": x, "field": processName(fieldsList[x])} for x in range(0, len(fieldsList))], 'nonEdit': nonEditNum}
    return render(request, 'admin/allVolInfo.html', context)




@login_required
@allowed_users(allowed_roles=['admin',])
def getVolDataArray(request):
    for i in volunteer.objects.all():
        i.save()

    data2 = list(volunteer.objects.values_list('dataTableInfo', flat=True))

    return JsonResponse(data2, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def updateVol(request):
    objId = request.POST.get('id')
    result = request.POST.getlist('result[]')
    test = volunteer.objects.get(id=objId)
    fieldsList = getKeys(test, exclude=set({'id','hoursentry'}))
    alertList = []

    for x in range(0, len(fieldsList)):
        try:
            #exec(f"test.{fieldsList[x]} = '{result[x]}'")
            setattr(test, fieldsList[x], result[x])
        except SyntaxError:
            cur = result[x].replace('\n', '\\n')
            #exec(f"test.{fieldsList[x]} = '{cur}'")
            setattr(test, fieldsList[x], cur)
        x += 1

    try:
        with transaction.atomic():
            test.save()
    except Exception as e:
        if type(e).__name__ == "ValidationError":
            theDict = ast.literal_eval(str(e))
            for singleField in theDict.keys():
                for mess in theDict[singleField]:
                    alertList.append(f"Field '{singleField}' - {mess}")
        elif type(e).__name__ == "ValueError":
            mess = str(e)
            field = mess[7:(mess[7:].find("'") + 7)]
            alertList.append(mess.replace(field, processName(field)))
        else: 
            alertList.append(str(e))
    #else:
        #rectifyGlobalVolHours()
        #updateGlobalVolHours()
    return JsonResponse(alertList, safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def deleteVol(request):
    entry = volunteer.objects.get(id=int(request.POST['id']))
    entry.delete()

    return JsonResponse("Deleted!", safe=False)




@login_required
@allowed_users(allowed_roles=['admin',])
def requestTrackInfo(request):
    return render(request, 'admin/requestTrackInfo.html')



from mainpg.models import ReqFormSettings

@login_required
@allowed_users(allowed_roles=['admin',])
def editRequestForm(request):
    if request.method == "POST":
        settingObj = ReqFormSettings.load()
        curSettings = json.loads(request.POST['list'])
        curLinks = json.loads(request.POST['curLinks'])
        alertList = set()
        groupSet = set()
        dbnSet = set()
        dspSet = set()
        dbnToDsp = {}
        for group in curSettings:
            if not group[0]:
                alertList.add(f"Group names cannot be blank")
            if group[0] in groupSet:
                alertList.add(f"Duplicate group name (\"{group[0]}\") is not allowed")
            else:
                groupSet.add(group[0])
            if not group[1]:
                alertList.add(f"Group \"{group[0]}\" - Group must have at least one field")
            
            for fieldInfo in group[1]:
                dbnToDsp[fieldInfo[0]] = fieldInfo[1]
                if not fieldInfo[0]:
                    alertList.add(f"Group \"{group[0]}\" contains a field with a blank database name")
                if fieldInfo[0] in dbnSet:
                    alertList.add(f"Duplicate database name (\"{fieldInfo[0]}\") is not allowed")
                else:
                    dbnSet.add(fieldInfo[0])
                
                if not fieldInfo[1]:
                    alertList.add(f"Group \"{group[0]}\" contains a field with a blank display name")
                if fieldInfo[1] in dspSet:
                    alertList.add(f"Duplicate display name (\"{fieldInfo[1]}\") is not allowed")
                else:
                    dspSet.add(fieldInfo[1])
                
                if fieldInfo[2] == 'dropdown':
                    if not fieldInfo[3]:
                        alertList.add(f"In group \"{group[0]}\", dropdown field \"{fieldInfo[0]}\" has no dropdown options")
                    elif '' in fieldInfo[3]:
                        alertList.add(f"In group \"{group[0]}\", dropdown field \"{fieldInfo[0]}\" has a blank dropdown option(s)")

                del fieldInfo[1]

        for field in curLinks:
            #curLinks[field] is a dictionary; its .keys() contains the database names of the linked fields
            ke = curLinks[field].keys()
            if '' in ke:
                alertList.add(f"Field with database name \"{field}\" has a blank link(s)")
            if not ke <= dbnSet:
                alertList.add(f"Field with database name \"{field}\" has a link to a field that does not exist")
            if field in ke:
                alertList.add(f"Field with database name \"{field}\" cannot have a link to itself")



        alertList = list(alertList)
        if alertList:
            return JsonResponse([False, alertList], safe=False)
        else:
            try:
                with transaction.atomic():
                    settingObj.fieldInfo = curSettings
                    settingObj.links = curLinks
                    settingObj.databaseDisplay = dbnToDsp
                    settingObj.save()
            except Exception as e:
                errList = []
                excepts(e, alertList, errList)
                return JsonResponse([False, alertList], safe=False)
            else:
                messages.success(request, "Form settings have been saved successfully.")
                return JsonResponse([True, "/edit-request-form"], safe=False)

    else:
        existingSet = ReqFormSettings.load()
        context = {
            'settingsDic': json.dumps(existingSet.fieldInfo),
            'linksDic': json.dumps(existingSet.links), 
            'databaseDisplay': json.dumps(existingSet.databaseDisplay), 
        }

        return render(request, "admin/editRequestForm.html", context)
